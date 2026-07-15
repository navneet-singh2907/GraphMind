"""Extensible parser registry for heterogeneous source documents."""

from __future__ import annotations

import csv
import json
from abc import ABC, abstractmethod
from html.parser import HTMLParser
from pathlib import Path
from typing import Iterable

from src.ingestion.models import ParsedDocument


class DocumentParser(ABC):
    """Adapter interface implemented by every source parser."""

    extensions: frozenset[str] = frozenset()

    def supports(self, path: Path) -> bool:
        return path.suffix.lower() in self.extensions

    @abstractmethod
    def parse(self, path: Path) -> list[ParsedDocument]:
        raise NotImplementedError


def _read_text(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1252"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="replace")


class TextParser(DocumentParser):
    extensions = frozenset({".txt", ".md", ".rst", ".log", ".yaml", ".yml"})

    def parse(self, path: Path) -> list[ParsedDocument]:
        return [ParsedDocument(title=path.stem, content=_read_text(path))]


class JsonParser(DocumentParser):
    extensions = frozenset({".json", ".jsonl"})

    def parse(self, path: Path) -> list[ParsedDocument]:
        raw = _read_text(path)
        if path.suffix.lower() == ".jsonl":
            rows = [json.loads(line) for line in raw.splitlines() if line.strip()]
            content = "\n".join(json.dumps(row, ensure_ascii=False, sort_keys=True) for row in rows)
        else:
            content = json.dumps(json.loads(raw), ensure_ascii=False, indent=2, sort_keys=True)
        return [ParsedDocument(title=path.stem, content=content)]


class CsvParser(DocumentParser):
    extensions = frozenset({".csv", ".tsv"})

    def parse(self, path: Path) -> list[ParsedDocument]:
        delimiter = "\t" if path.suffix.lower() == ".tsv" else ","
        with path.open("r", encoding="utf-8-sig", newline="") as handle:
            rows = list(csv.DictReader(handle, delimiter=delimiter))
        content = "\n".join(
            "; ".join(f"{key}: {value}" for key, value in row.items()) for row in rows
        )
        return [ParsedDocument(title=path.stem, content=content, metadata={"rows": len(rows)})]


class PdfParser(DocumentParser):
    extensions = frozenset({".pdf"})

    def parse(self, path: Path) -> list[ParsedDocument]:
        from pypdf import PdfReader

        documents = []
        for page_number, page in enumerate(PdfReader(str(path)).pages, start=1):
            content = (page.extract_text() or "").strip()
            if content:
                documents.append(
                    ParsedDocument(
                        title=path.stem,
                        content=content,
                        metadata={"page": page_number},
                    )
                )
        return documents


class DocxParser(DocumentParser):
    extensions = frozenset({".docx"})

    def parse(self, path: Path) -> list[ParsedDocument]:
        from docx import Document

        document = Document(str(path))
        blocks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells]
                if any(values):
                    blocks.append(" | ".join(values))
        return [ParsedDocument(title=path.stem, content="\n\n".join(blocks))]


class XlsxParser(DocumentParser):
    extensions = frozenset({".xlsx", ".xlsm"})

    def parse(self, path: Path) -> list[ParsedDocument]:
        from openpyxl import load_workbook

        workbook = load_workbook(path, read_only=True, data_only=True)
        documents = []
        try:
            for sheet in workbook.worksheets:
                rows = []
                for values in sheet.iter_rows(values_only=True):
                    cells = ["" if value is None else str(value) for value in values]
                    if any(cells):
                        rows.append(" | ".join(cells))
                if rows:
                    documents.append(
                        ParsedDocument(
                            title=f"{path.stem} — {sheet.title}",
                            content="\n".join(rows),
                            metadata={"sheet": sheet.title, "rows": len(rows)},
                        )
                    )
        finally:
            workbook.close()
        return documents


class _TextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []
        self._ignored_depth = 0

    def handle_starttag(self, tag: str, attrs) -> None:
        if tag in {"script", "style", "noscript"}:
            self._ignored_depth += 1

    def handle_endtag(self, tag: str) -> None:
        if tag in {"script", "style", "noscript"} and self._ignored_depth:
            self._ignored_depth -= 1

    def handle_data(self, data: str) -> None:
        if not self._ignored_depth and data.strip():
            self.parts.append(data.strip())


class HtmlParser(DocumentParser):
    extensions = frozenset({".html", ".htm"})

    def parse(self, path: Path) -> list[ParsedDocument]:
        extractor = _TextExtractor()
        extractor.feed(_read_text(path))
        return [ParsedDocument(title=path.stem, content="\n".join(extractor.parts))]


class ParserRegistry:
    def __init__(self, parsers: Iterable[DocumentParser] | None = None) -> None:
        self.parsers = list(
            parsers
            or [TextParser(), JsonParser(), CsvParser(), PdfParser(), DocxParser(), XlsxParser(), HtmlParser()]
        )

    @property
    def supported_extensions(self) -> set[str]:
        return {extension for parser in self.parsers for extension in parser.extensions}

    def parser_for(self, path: Path) -> DocumentParser:
        for parser in self.parsers:
            if parser.supports(path):
                return parser
        raise ValueError(f"No parser is registered for {path.suffix or path.name}")

    def parse(self, path: Path) -> list[ParsedDocument]:
        return self.parser_for(path).parse(path)
